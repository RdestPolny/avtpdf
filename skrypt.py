"""
Redaktor AI - Interaktywny Procesor Dokumentów
===============================================

INSTALACJA ZALEŻNOŚCI:
----------------------
pip uninstall docx  # WAŻNE: Usuń złą bibliotekę jeśli jest zainstalowana!
pip install streamlit PyMuPDF openai python-docx mammoth

MINIMALNA KONFIGURACJA (tylko PDF):
pip install streamlit PyMuPDF openai

URUCHOMIENIE:
streamlit run skrypt.py
"""

import streamlit as st
import fitz  # PyMuPDF
from openai import AsyncOpenAI
import io
import zipfile
import json
from pathlib import Path
import re
import asyncio
from typing import List, Dict, Optional
from dataclasses import dataclass

# Opcjonalne importy - aplikacja będzie działać bez nich (tylko PDF)
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

# ===== KONFIGURACJA =====

PROJECTS_DIR = Path("pdf_processor_projects")
BATCH_SIZE = 10
MAX_RETRIES = 3
DEFAULT_MODEL = 'gpt-4o-mini'

SESSION_STATE_DEFAULTS = {
    'processing_status': 'idle',
    'document': None,
    'current_page': 0,
    'total_pages': 0,
    'extracted_pages': [],
    'project_name': None,
    'next_batch_start_index': 0,
    'uploaded_filename': None,
    'api_key': None,
    'model': DEFAULT_MODEL,
    'meta_tags': {},
    'seo_articles': {},
    'project_loaded_and_waiting_for_file': False,
    'processing_mode': 'all',
    'start_page': 1,
    'end_page': 1,
    'processing_end_page_index': 0,
    'article_page_groups_input': '',
    'article_groups': [],
    'next_article_index': 0,
    'file_type': None
}

# ===== KLASY POMOCNICZE =====

@dataclass
class PageContent:
    """Reprezentuje zawartość pojedynczej strony"""
    page_number: int
    text: str
    images: List[Dict] = None
    
    def __post_init__(self):
        if self.images is None:
            self.images = []

class DocumentHandler:
    """Klasa do obsługi różnych formatów dokumentów"""
    
    def __init__(self, file_bytes: bytes, filename: str):
        self.file_bytes = file_bytes
        self.filename = filename
        self.file_type = self._detect_file_type(filename)
        self._document = None
        self._html_content = None
        self._load_document()
    
    def _detect_file_type(self, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            if not DOCX_AVAILABLE:
                raise ValueError("Format DOCX nie jest obsługiwany. Zainstaluj: pip install python-docx")
            return 'docx'
        elif ext == '.doc':
            if not MAMMOTH_AVAILABLE:
                raise ValueError("Format DOC nie jest obsługiwany. Zainstaluj: pip install mammoth")
            return 'doc'
        else:
            raise ValueError(f"Nieobsługiwany format pliku: {ext}")
    
    def _load_document(self):
        """Ładuje dokument w odpowiednim formacie"""
        if self.file_type == 'pdf':
            self._document = fitz.open(stream=self.file_bytes, filetype="pdf")
        elif self.file_type == 'docx':
            self._document = DocxDocument(io.BytesIO(self.file_bytes))
        elif self.file_type == 'doc':
            result = mammoth.convert_to_html(io.BytesIO(self.file_bytes))
            self._html_content = result.value
            self._document = None
    
    def get_page_count(self) -> int:
        """Zwraca liczbę stron w dokumencie"""
        if self.file_type == 'pdf':
            return len(self._document)
        elif self.file_type == 'docx':
            all_text = '\n\n'.join([p.text for p in self._document.paragraphs])
            words = all_text.split()
            return max(1, len(words) // 500 + (1 if len(words) % 500 > 0 else 0))
        elif self.file_type == 'doc':
            words = self._html_content.split()
            return max(1, len(words) // 500 + (1 if len(words) % 500 > 0 else 0))
        return 0
    
    def get_page_content(self, page_index: int) -> PageContent:
        """Zwraca zawartość pojedynczej strony"""
        if self.file_type == 'pdf':
            page = self._document.load_page(page_index)
            text = page.get_text("text")
            images = self._extract_images_from_pdf_page(page_index)
            return PageContent(page_index + 1, text, images)
        elif self.file_type == 'docx':
            return self._get_docx_page_content(page_index)
        elif self.file_type == 'doc':
            return self._get_doc_page_content(page_index)
    
    def _get_docx_page_content(self, page_index: int) -> PageContent:
        """Pobiera zawartość z DOCX - dzieli na fragmenty"""
        all_paragraphs = self._document.paragraphs
        words_per_page = 500
        
        all_text = '\n\n'.join([p.text for p in all_paragraphs])
        words = all_text.split()
        
        start_word = page_index * words_per_page
        end_word = min(start_word + words_per_page, len(words))
        
        page_text = ' '.join(words[start_word:end_word])
        images = self._extract_images_from_docx()
        
        return PageContent(page_index + 1, page_text, images)
    
    def _get_doc_page_content(self, page_index: int) -> PageContent:
        """Pobiera zawartość z DOC (HTML)"""
        text = re.sub('<[^<]+?>', '', self._html_content)
        words = text.split()
        words_per_page = 500
        
        start_word = page_index * words_per_page
        end_word = min(start_word + words_per_page, len(words))
        
        page_text = ' '.join(words[start_word:end_word])
        
        return PageContent(page_index + 1, page_text, [])
    
    def _extract_images_from_pdf_page(self, page_index: int) -> List[Dict]:
        """Wyciąga obrazy z strony PDF"""
        images = []
        if self.file_type != 'pdf':
            return images
        
        try:
            page = self._document.load_page(page_index)
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = self._document.extract_image(xref)
                if base_image and base_image.get("width", 0) > 100 and base_image.get("height", 0) > 100:
                    images.append({
                        'image': base_image['image'],
                        'ext': base_image['ext'],
                        'index': img_index
                    })
        except Exception as e:
            st.warning(f"Nie udało się wyekstraktować obrazów ze strony {page_index + 1}: {e}")
        
        return images
    
    def _extract_images_from_docx(self) -> List[Dict]:
        """Wyciąga obrazy z dokumentu DOCX"""
        images = []
        try:
            for rel in self._document.part.rels.values():
                if "image" in rel.target_ref:
                    img_data = rel.target_part.blob
                    ext = rel.target_ref.split('.')[-1]
                    images.append({
                        'image': img_data,
                        'ext': ext,
                        'index': len(images)
                    })
        except Exception as e:
            st.warning(f"Nie udało się wyekstraktować obrazów z DOCX: {e}")
        
        return images
    
    def render_page_as_image(self, page_index: int) -> Optional[bytes]:
        """Renderuje stronę jako obraz (tylko dla PDF)"""
        if self.file_type != 'pdf':
            return None
        
        try:
            page = self._document.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            return pix.tobytes("png")
        except Exception as e:
            st.error(f"Błąd podczas renderowania strony {page_index + 1}: {e}")
            return None

# ===== LOGIKA AI =====

class AIProcessor:
    """Klasa obsługująca komunikację z OpenAI API"""
    
    def __init__(self, api_key: str, model: str = DEFAULT_MODEL):
        self.client = AsyncOpenAI(api_key=api_key)
        self.model = model
    
    def get_system_prompt(self) -> str:
        """Zwraca system prompt dla przetwarzania artykułów"""
        return """Jesteś precyzyjnym asystentem redakcyjnym. Twoim celem jest przekształcenie surowego tekstu w czytelny, dobrze zorganizowany artykuł internetowy.

ZASADA NADRZĘDNA: WIERNOŚĆ TREŚCI, ELASTYCZNOŚĆ FORMY.
- Nie zmieniaj oryginalnych sformułowań ani nie parafrazuj tekstu (chyba że to konieczne dla czytelności). Przenieś treść 1:1.
- Twoja rola polega na dodawaniu elementów strukturalnych i czyszczeniu śmieci.

INSTRUKCJE SPECJALNE (KRYTYCZNE):
1. **LISTY:** Jeśli widzisz wyliczenia (punkty, myślniki), formatuj je jako standardową listę Markdown:
   - Element listy 1
   - Element listy 2
2. **PRZYPISY/INDEKSY:** Jeśli wykryjesz indeksy przypisów (małe cyfry na końcu zdań lub słów), formatuj je używając tagu HTML: `<sup>1</sup>`, `<sup>2</sup>`.
3. **USUWANIE PODPISÓW I ŚMIECI:** BEZWZGLĘDNIE USUWAJ:
   - Podpisy pod zdjęciami (np. "Rys. 1. Widok...", "Fot. Jan Kowalski").
   - Źródła grafik i tabel (np. "Źródło: opracowanie własne").
   - Numery stron, nagłówki i stopki redakcyjne.
   - Etykiety typu "NEWS FLASH".

DOZWOLONE MODYFIKACJE STRUKTURALNE:
1. Tytuł Główny: `# Tytuł`
2. Śródtytuły: `## Śródtytuł` (używaj ich do rozbijania 'ściany tekstu').
3. Pogrubienia: `**tekst**` (dla kluczowych terminów i nazw własnych).
4. Podział na sekcje: `---` (jeśli na stronie są dwa niepowiązane tematy).

WYMAGANIA KRYTYCZNE:
- Twoja odpowiedź musi być WYŁĄCZNIE i BEZWZGLĘDNIE poprawnym obiektem JSON.
- NIE używaj markdown code blocks (```json). Zwróć TYLKO czysty JSON.

FORMAT ODPOWIEDZI:
{"type": "ARTYKUŁ" lub "REKLAMA", "formatted_text": "Sformatowany tekst w markdown."}"""
    
    def get_meta_tags_prompt(self) -> str:
        """Zwraca prompt dla generowania meta tagów"""
        return """Jesteś ekspertem SEO. Na podstawie poniższego tekstu artykułu, wygeneruj chwytliwy meta title i zwięzły meta description.

WYMAGANIA:
- Meta title: max 60 znaków.
- Meta description: max 160 znaków.
- Odpowiedź zwróć jako czysty obiekt JSON bez markdown code blocks.

FORMAT ODPOWIEDZI:
{"meta_title": "Tytuł meta", "meta_description": "Opis meta."}"""

    def get_seo_prompt(self) -> str:
        """Zwraca prompt dla optymalizacji artykułu pod kątem SEO."""
        return """Jesteś światowej klasy strategiem SEO i copywriterem. Twoim zadaniem jest przepisanie dostarczonego artykułu, aby był maksymalnie zoptymalizowany pod kątem wyszukiwarek i angażujący dla czytelników online.

ZASADY KRYTYCZNE:
1.  **WIERNOŚĆ FAKTÓW**: Musisz bazować WYŁĄCZNIE na informacjach zawartych w oryginalnym tekście. Nie dodawaj żadnych nowych faktów, danych ani opinii. Twoja rola to restrukturyzacja i optymalizacja.
2.  **ODWRÓCONA PIRAMIDA**: Zastosuj zasadę odwróconej piramidy. Najważniejsze informacje, kluczowe wnioski i odpowiedzi na potencjalne pytania czytelnika umieść na samym początku artykułu.
3.  **STRUKTURA I CZYTELNOŚĆ**:
    *   Stwórz nowy, chwytliwy tytuł zoptymalizowany pod kątem potencjalnych fraz kluczowych (H1).
    *   Podziel tekst na logiczne sekcje za pomocą śródtytułów (H2, H3).
    *   Używaj list punktowanych, jeśli to możliwe, aby zwiększyć czytelność.
    *   Stosuj pogrubienia (`**tekst**`) dla najważniejszych terminów.
4.  **JĘZYK**: Używaj aktywnego, dynamicznego języka. Unikaj strony biernej. Pisz bezpośrednio do czytelnika.

WYMAGANIA FORMATOWANIA:
- Twoja odpowiedź musi być WYŁĄCZNIE poprawnym obiektem JSON.
- NIE używaj bloków kodu markdown (```json).

FORMAT ODPOWIEDZI JSON:
{"seo_title": "Nowy, zoptymalizowany pod SEO tytuł artykułu", "seo_article_markdown": "Pełna treść przepisanego artykułu w formacie Markdown, z nagłówkami, listami i pogrubieniami."}"""

    async def process_text(self, text: str, system_prompt: str, max_tokens: int = 4096) -> Dict:
        """Przetwarza tekst przez OpenAI API"""
        last_error = None
        content = ""
        
        for attempt in range(MAX_RETRIES):
            try:
                response = await self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": text}
                    ],
                    max_tokens=max_tokens,
                    temperature=0.3,
                    response_format={"type": "json_object"}
                )
                
                content = response.choices[0].message.content
                
                if not content:
                    raise ValueError("API zwróciło pustą odpowiedź.")
                
                return json.loads(content)
                
            except json.JSONDecodeError as e:
                last_error = e
                st.warning(f"Próba {attempt + 1}: Błąd dekodowania JSON. Ponawiam próbę...")
                await asyncio.sleep(1)
                continue
            except Exception as e:
                last_error = e
                st.warning(f"Próba {attempt + 1}: Wystąpił błąd API: {e}. Ponawiam próbę...")
                await asyncio.sleep(1)
                continue
        
        return {
            "error": f"Błąd po {MAX_RETRIES} próbach.",
            "last_known_error": str(last_error),
            "raw_response": content
        }
    
    async def process_page(self, page_content: PageContent) -> Dict:
        """Przetwarza pojedynczą stronę"""
        page_data = {"page_number": page_content.page_number}
        
        if len(page_content.text.split()) < 20:
            page_data["type"] = "pominięta"
            page_data["formatted_content"] = "<i>Strona zawiera zbyt mało tekstu.</i>"
            return page_data
        
        result = await self.process_text(page_content.text, self.get_system_prompt(), max_tokens=4096)
        
        if "error" in result:
            page_data["type"] = "błąd"
            page_data["formatted_content"] = f"""<div class="error-box">
                <strong>{result['error']}</strong><br>
                <i>Ostatni błąd: {result['last_known_error']}</i><br>
                <details><summary>Pokaż surową odpowiedź</summary><pre>{result['raw_response']}</pre></details>
            </div>"""
        else:
            page_data["type"] = result.get("type", "nieznany").lower()
            formatted_text = result.get("formatted_text", "")
            
            if page_data["type"] == "artykuł":
                page_data["formatted_content"] = markdown_to_html(formatted_text)
                page_data["raw_markdown"] = formatted_text
            else:
                page_data["formatted_content"] = f"<i>Zidentyfikowano jako: <strong>{page_data['type'].upper()}</strong>.</i>"
        
        return page_data
    
    async def process_article_group(self, pages_content: List[PageContent]) -> Dict:
        """Przetwarza grupę stron jako jeden artykuł"""
        page_numbers = [p.page_number for p in pages_content]
        
        combined_text = "\n\n".join([
            f"--- STRONA {p.page_number} ---\n{p.text.strip()}"
            for p in pages_content
        ])
        
        result = await self.process_text(combined_text, self.get_system_prompt(), max_tokens=8192)
        
        article_data = {"page_numbers": page_numbers}
        
        if "error" in result:
            article_data["type"] = "błąd"
            article_data["formatted_content"] = f"""<div class='error-box'>
                <strong>{result['error']}</strong><br>
                <i>Ostatni błąd: {result['last_known_error']}</i><br>
                <details><summary>Pokaż surową odpowiedź</summary><pre>{result['raw_response']}</pre></details>
            </div>"""
        else:
            article_data["type"] = result.get("type", "nieznany").lower()
            formatted_text = result.get("formatted_text", "")
            
            if article_data["type"] == "artykuł":
                article_data["formatted_content"] = markdown_to_html(formatted_text)
                article_data["raw_markdown"] = formatted_text
            else:
                article_data["formatted_content"] = f"<i>Zidentyfikowano jako: <strong>{article_data['type'].upper()}</strong>.</i>"
        
        return article_data
    
    async def generate_meta_tags(self, article_text: str) -> Dict:
        """Generuje meta tagi dla artykułu"""
        prompt = self.get_meta_tags_prompt()
        return await self.process_text(article_text[:4000], prompt, max_tokens=200)

    async def generate_seo_article(self, article_text: str) -> Dict:
        """Przepisuje artykuł pod kątem SEO"""
        prompt = self.get_seo_prompt()
        return await self.process_text(article_text, prompt, max_tokens=4096)

# ===== FUNKCJE POMOCNICZE =====

def markdown_to_html(text: str) -> str:
    """Konwertuje markdown na HTML z obsługą list i indeksów"""
    # 1. Obsługa struktury nagłówków i linii
    text = text.replace('\n---\n', '\n<hr>\n')
    text = re.sub(r'^\s*# (.*?)\s*$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*## (.*?)\s*$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*### (.*?)\s*$', r'<h4>\1</h4>', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
    
    # 2. Obsługa list (zamiana linii z myślnikami na <ul><li>...</li></ul>)
    lines = text.split('\n')
    new_lines = []
    in_list = False
    
    for line in lines:
        stripped = line.strip()
        # Wykrywanie elementu listy (myślnik lub gwiazdka na początku)
        if stripped.startswith(('- ', '* ')):
            if not in_list:
                new_lines.append("<ul>")
                in_list = True
            content = stripped[2:] # Usunięcie "- "
            new_lines.append(f"<li>{content}</li>")
        else:
            if in_list:
                new_lines.append("</ul>")
                in_list = False
            new_lines.append(line)
            
    if in_list:
        new_lines.append("</ul>")
        
    text = '\n'.join(new_lines)
    
    # 3. Podział na akapity (z pominięciem bloków, które już są HTMLem)
    paragraphs = text.split('\n\n')
    html_content = []
    
    for para in paragraphs:
        stripped_para = para.strip()
        if not stripped_para:
            continue
            
        # Sprawdź czy to element blokowy HTML (nagłówek, lista, hr)
        if stripped_para.startswith(('<h', '<hr', '<ul', '<li')):
            html_content.append(stripped_para)
        else:
            # Zachowaj tagi <sup> jeśli są w tekście, zamień nową linię na <br>
            formatted_para = stripped_para.replace(chr(10), '<br>')
            html_content.append(f"<p>{formatted_para}</p>")
    
    return ''.join(html_content)

def markdown_to_clean_html(markdown_text: str, page_number: int = None) -> str:
    """
    Konwertuje markdown na czysty HTML bez stylowania
    Tylko struktura: h1-h4, p, strong, hr, ul, li, sup
    """
    html = markdown_text
    
    html = html.replace('\n---\n', '\n<hr>\n')
    html = html.replace('\n--- \n', '\n<hr>\n')
    
    html = re.sub(r'^\s*# (.*?)\s*$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    html = re.sub(r'^\s*## (.*?)\s*$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^\s*### (.*?)\s*$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^\s*#### (.*?)\s*$', r'<h4>\1</h4>', html, flags=re.MULTILINE)
    html = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html)

    # Obsługa list w czystym HTML
    lines = html.split('\n')
    new_lines = []
    in_list = False
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith(('- ', '* ')):
            if not in_list:
                new_lines.append("<ul>")
                in_list = True
            content = stripped[2:]
            new_lines.append(f"<li>{content}</li>")
        else:
            if in_list:
                new_lines.append("</ul>")
                in_list = False
            new_lines.append(line)
    
    if in_list:
        new_lines.append("</ul>")
        
    html = '\n'.join(new_lines)
    
    # Podział na akapity
    paragraphs = html.split('\n\n')
    formatted_paragraphs = []
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # Jeśli linia zaczyna się od tagu HTML, nie pakuj jej w <p>
        if para.startswith(('<h1', '<h2', '<h3', '<h4', '<hr', '<p', '<ul')):
            formatted_paragraphs.append(para)
        else:
            para_with_breaks = para.replace('\n', '<br>\n')
            formatted_paragraphs.append(f'<p>{para_with_breaks}</p>')
    
    return '\n'.join(formatted_paragraphs)

def generate_full_html_document(content: str, title: str = "Artykuł", meta_title: str = None, meta_description: str = None) -> str:
    """
    Generuje pełny dokument HTML z czystą strukturą (bez CSS)
    """
    meta_tags = ""
    if meta_title:
        meta_tags += f'    <meta name="title" content="{meta_title}">\n'
    if meta_description:
        meta_tags += f'    <meta name="description" content="{meta_description}">\n'
    
    html_doc = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
{meta_tags}</head>
<body>
{content}
</body>
</html>"""
    
    return html_doc

def get_article_html_from_page(page_index: int) -> Optional[Dict]:
    """
    Pobiera czysty HTML artykułu dla danej strony
    """
    page_result = st.session_state.extracted_pages[page_index]
    
    if not page_result or page_result.get('type') != 'artykuł':
        return None
    
    if 'raw_markdown' not in page_result:
        return None
    
    group_pages = page_result.get('group_pages', [])
    
    if group_pages and len(group_pages) > 1:
        first_page_index = group_pages[0] - 1
        first_page_result = st.session_state.extracted_pages[first_page_index]
        
        markdown_content = first_page_result.get('raw_markdown', '')
        title = f"Artykuł ze stron {group_pages[0]}-{group_pages[-1]}"
        pages = group_pages
    else:
        markdown_content = page_result.get('raw_markdown', '')
        title = f"Artykuł ze strony {page_index + 1}"
        pages = [page_index + 1]
    
    html_content = markdown_to_clean_html(markdown_content)
    
    meta_title = None
    meta_description = None
    
    if page_index in st.session_state.meta_tags:
        tags = st.session_state.meta_tags[page_index]
        if 'error' not in tags:
            meta_title = tags.get('meta_title')
            meta_description = tags.get('meta_description')
    
    html_document = generate_full_html_document(
        html_content,
        title=title,
        meta_title=meta_title,
        meta_description=meta_description
    )
    
    return {
        'html_content': html_content,
        'html_document': html_document,
        'title': title,
        'pages': pages,
        'meta_title': meta_title,
        'meta_description': meta_description
    }

def sanitize_filename(name: str) -> str:
    """Sanityzuje nazwę pliku"""
    if not name:
        return "unnamed_project"
    sanitized = re.sub(r'[\\/*?:"<>|]', "_", str(name))
    return re.sub(r'_{2,}', "_", sanitized).strip("_") or "unnamed_project"

def create_zip_archive(data: List[Dict]) -> bytes:
    """Tworzy archiwum ZIP"""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in data:
            zf.writestr(item['name'], item['content'])
    return zip_buffer.getvalue()

def parse_page_groups(input_text: str, total_pages: int) -> List[List[int]]:
    """Parsuje zakresy stron z tekstu wejściowego"""
    if not input_text:
        raise ValueError("Nie podano zakresów stron.")
    
    groups = []
    used_pages = set()
    
    for line in re.split(r'[;\n]+', input_text):
        line = line.strip()
        if not line:
            continue
        
        pages = []
        for part in re.split(r'[;,]+', line):
            part = part.strip()
            if not part:
                continue
            
            if '-' in part:
                start_str, end_str = part.split('-', 1)
                if not start_str.isdigit() or not end_str.isdigit():
                    raise ValueError(f"Niepoprawny zakres stron: '{part}'.")
                
                start, end = int(start_str), int(end_str)
                if start > end:
                    raise ValueError(f"Zakres stron musi być rosnący: '{part}'.")
                if start < 1 or end > total_pages:
                    raise ValueError(f"Zakres '{part}' wykracza poza liczbę stron dokumentu.")
                
                pages.extend(range(start, end + 1))
            else:
                if not part.isdigit():
                    raise ValueError(f"Niepoprawny numer strony: '{part}'.")
                
                page = int(part)
                if page < 1 or page > total_pages:
                    raise ValueError(f"Strona '{page}' wykracza poza dokument.")
                
                pages.append(page)
        
        if not pages:
            continue
        
        pages = sorted(dict.fromkeys(pages))
        
        if any(p in used_pages for p in pages):
            raise ValueError(f"Strony {pages} zostały już przypisane do innego artykułu.")
        
        used_pages.update(pages)
        groups.append(pages)
    
    if not groups:
        raise ValueError("Nie znaleziono żadnych poprawnych zakresów stron.")
    
    return groups

# ===== ZARZĄDZANIE PROJEKTAMI =====

def ensure_projects_dir() -> bool:
    """Tworzy katalog projektów jeśli nie istnieje"""
    try:
        PROJECTS_DIR.mkdir(exist_ok=True)
        return True
    except Exception as e:
        st.error(f"Nie można utworzyć katalogu projektów: {e}")
        return False

def get_existing_projects() -> List[str]:
    """Zwraca listę istniejących projektów"""
    if not ensure_projects_dir():
        return []
    return [d.name for d in PROJECTS_DIR.iterdir() if d.is_dir()]

def save_project():
    """Zapisuje projekt do pliku"""
    if not st.session_state.project_name or not ensure_projects_dir():
        st.error("Nie można zapisać projektu: brak nazwy projektu.")
        return
    
    project_path = PROJECTS_DIR / st.session_state.project_name
    project_path.mkdir(exist_ok=True)
    
    state_to_save = {
        k: v for k, v in st.session_state.items()
        if k not in ['document', 'project_loaded_and_waiting_for_file']
    }
    state_to_save['extracted_pages'] = [
        p for p in st.session_state.extracted_pages if p is not None
    ]
    
    try:
        with open(project_path / "project_state.json", "w", encoding="utf-8") as f:
            json.dump(state_to_save, f, indent=2, ensure_ascii=False)
        st.toast(f"✅ Projekt '{st.session_state.project_name}' został zapisany!", icon="💾")
    except Exception as e:
        st.error(f"Błąd podczas zapisywania projektu: {e}")

def load_project(project_name: str):
    """Ładuje projekt z pliku"""
    project_file = PROJECTS_DIR / project_name / "project_state.json"
    
    if not project_file.exists():
        st.error(f"Plik projektu '{project_name}' nie istnieje.")
        return
    
    try:
        with open(project_file, "r", encoding="utf-8") as f:
            state_to_load = json.load(f)
        
        for key, value in state_to_load.items():
            if key != 'document':
                st.session_state[key] = value
        
        total_pages = st.session_state.get('total_pages', 0)
        st.session_state.extracted_pages = [None] * total_pages
        
        for page_data in state_to_load.get('extracted_pages', []):
            page_num_one_based = page_data.get('page_number')
            if page_num_one_based and 1 <= page_num_one_based <= total_pages:
                st.session_state.extracted_pages[page_num_one_based - 1] = page_data
        
        st.session_state.document = None
        st.session_state.project_loaded_and_waiting_for_file = True
        
        st.success(f"✅ Załadowano projekt '{project_name}'. Wgraj powiązany plik, aby kontynuować.")
    except Exception as e:
        st.error(f"Błąd podczas ładowania projektu: {e}")

# ===== OBSŁUGA PLIKÓW =====

def handle_file_upload(uploaded_file):
    """Obsługuje wgranie pliku"""
    try:
        with st.spinner("Ładowanie pliku..."):
            file_bytes = uploaded_file.read()
            document = DocumentHandler(file_bytes, uploaded_file.name)
            
            if st.session_state.project_loaded_and_waiting_for_file:
                if document.get_page_count() != st.session_state.total_pages:
                    st.error(
                        f"Błąd: Wgrany plik ma {document.get_page_count()} stron, "
                        f"a projekt oczekuje {st.session_state.total_pages}. Wgraj właściwy plik."
                    )
                    return
                
                st.session_state.document = document
                st.session_state.uploaded_filename = uploaded_file.name
                st.session_state.file_type = document.file_type
                st.session_state.project_loaded_and_waiting_for_file = False
                st.success("✅ Plik pomyślnie dopasowany do projektu.")
            else:
                for key, value in SESSION_STATE_DEFAULTS.items():
                    if key != 'api_key':
                        st.session_state[key] = value
                
                st.session_state.document = document
                st.session_state.uploaded_filename = uploaded_file.name
                st.session_state.file_type = document.file_type
                st.session_state.project_name = sanitize_filename(Path(uploaded_file.name).stem)
                st.session_state.total_pages = document.get_page_count()
                st.session_state.extracted_pages = [None] * document.get_page_count()
                st.session_state.end_page = document.get_page_count()
                
                st.success(f"✅ Załadowano plik: {uploaded_file.name} ({document.file_type.upper()})")
    
    except Exception as e:
        st.error(f"❌ Błąd ładowania pliku: {e}")
        st.session_state.document = None
    
    st.rerun()

# ===== PRZETWARZANIE AI =====

async def process_batch(ai_processor: AIProcessor, start_index: int):
    """Przetwarza batch stron"""
    processing_limit = st.session_state.processing_end_page_index + 1
    end_index = min(start_index + BATCH_SIZE, processing_limit)
    
    tasks = []
    for i in range(start_index, end_index):
        if st.session_state.document:
            page_content = st.session_state.document.get_page_content(i)
            tasks.append(ai_processor.process_page(page_content))
    
    results = await asyncio.gather(*tasks, return_exceptions=True)
    
    for i, result in enumerate(results):
        page_index = start_index + i
        if isinstance(result, Exception):
            st.session_state.extracted_pages[page_index] = {
                "page_number": page_index + 1,
                "type": "błąd",
                "formatted_content": f"Błąd: {result}"
            }
        else:
            st.session_state.extracted_pages[page_index] = result

def start_ai_processing():
    """Rozpoczyna przetwarzanie AI"""
    if st.session_state.processing_mode == 'article':
        try:
            groups = parse_page_groups(
                st.session_state.article_page_groups_input,
                st.session_state.total_pages
            )
            
            for group in groups:
                for page in group:
                    st.session_state.extracted_pages[page - 1] = None
            
            st.session_state.article_groups = groups
            st.session_state.next_article_index = 0
            st.session_state.processing_status = 'in_progress'
            
            # Automatycznie przejdź do pierwszej strony pierwszego artykułu
            if groups:
                st.session_state.current_page = groups[0][0] - 1
            
        except ValueError as e:
            st.error(str(e))
            return
    else:
        if st.session_state.processing_mode == 'all':
            start_idx = 0
            end_idx = st.session_state.total_pages - 1
        else:
            start_idx = st.session_state.start_page - 1
            end_idx = st.session_state.end_page - 1
        
        if start_idx > end_idx:
            st.error("Strona początkowa nie może być większa niż końcowa.")
            return
        
        for i in range(start_idx, end_idx + 1):
            st.session_state.extracted_pages[i] = None
        
        st.session_state.processing_status = 'in_progress'
        st.session_state.next_batch_start_index = start_idx
        st.session_state.processing_end_page_index = end_idx
        
        # Automatycznie przejdź do pierwszej przetwarzanej strony
        st.session_state.current_page = start_idx

def run_ai_processing_loop():
    """Główna pętla przetwarzania AI"""
    if not st.session_state.api_key:
        st.error("Klucz API OpenAI nie jest skonfigurowany.")
        st.session_state.processing_status = 'idle'
        return
    
    ai_processor = AIProcessor(st.session_state.api_key, st.session_state.model)
    
    if st.session_state.processing_mode == 'article':
        if st.session_state.next_article_index < len(st.session_state.article_groups):
            article_pages = st.session_state.article_groups[st.session_state.next_article_index]
            
            pages_content = []
            for page_num in article_pages:
                if st.session_state.document and 0 <= page_num - 1 < st.session_state.total_pages:
                    pages_content.append(
                        st.session_state.document.get_page_content(page_num - 1)
                    )
            
            article_result = asyncio.run(
                ai_processor.process_article_group(pages_content)
            )
            
            for page in article_pages:
                page_index = page - 1
                if 0 <= page_index < len(st.session_state.extracted_pages):
                    entry = {
                        key: value for key, value in article_result.items()
                        if key != 'page_numbers'
                    }
                    entry['page_number'] = page
                    entry['group_pages'] = article_pages
                    entry['is_group_lead'] = (page == article_pages[0])
                    st.session_state.extracted_pages[page_index] = entry
            
            st.session_state.next_article_index += 1
        else:
            st.session_state.processing_status = 'complete'
    else:
        if st.session_state.next_batch_start_index <= st.session_state.processing_end_page_index:
            asyncio.run(
                process_batch(ai_processor, st.session_state.next_batch_start_index)
            )
            st.session_state.next_batch_start_index += BATCH_SIZE
        else:
            st.session_state.processing_status = 'complete'
    
    st.rerun()

# ===== UI COMPONENTS =====

def init_session_state():
    """Inicjalizuje stan sesji"""
    if 'api_key' not in st.session_state or st.session_state.api_key is None:
        st.session_state.api_key = st.secrets.get("openai", {}).get("api_key")
    
    for key, value in SESSION_STATE_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = value

def render_sidebar():
    """Renderuje panel boczny"""
    with st.sidebar:
        st.header("⚙️ Konfiguracja Projektu")
        
        projects = get_existing_projects()
        selected_project = st.selectbox(
            "Wybierz istniejący projekt",
            ["Nowy projekt"] + projects
        )
        
        if st.button("Załaduj projekt", disabled=(selected_project == "Nowy projekt")):
            load_project(selected_project)
            st.rerun()
        
        st.divider()
        
        supported_formats = ["pdf"]
        if DOCX_AVAILABLE:
            supported_formats.append("docx")
        if MAMMOTH_AVAILABLE:
            supported_formats.append("doc")
        
        file_label = f"Wybierz plik ({', '.join(f.upper() for f in supported_formats)})"
        
        uploaded_file = st.file_uploader(
            file_label,
            type=supported_formats
        )
        
        if uploaded_file:
            if (st.session_state.project_loaded_and_waiting_for_file or
                uploaded_file.name != st.session_state.get('uploaded_filename')):
                handle_file_upload(uploaded_file)
        
        if st.session_state.document:
            st.divider()
            st.subheader("🤖 Opcje Przetwarzania")
            
            st.radio(
                "Wybierz tryb:",
                ('all', 'range', 'article'),
                captions=[
                    "Cały dokument (strona po stronie)",
                    "Zakres stron (strona po stronie)",
                    "Artykuł wielostronicowy (jedno zapytanie)"
                ],
                key='processing_mode',
                horizontal=False
            )
            
            if st.session_state.processing_mode == 'range':
                c1, c2 = st.columns(2)
                c1.number_input(
                    "Od strony",
                    min_value=1,
                    max_value=st.session_state.total_pages,
                    key='start_page'
                )
                c2.number_input(
                    "Do strony",
                    min_value=st.session_state.start_page,
                    max_value=st.session_state.total_pages,
                    key='end_page'
                )
            
            elif st.session_state.processing_mode == 'article':
                st.info("Podaj grupy stron dla artykułów wielostronicowych. Każda grupa zostanie przetworzona w jednym zapytaniu do AI.")
                st.text_area(
                    "Zakresy stron artykułów (np. 1-3; 5,6)",
                    key='article_page_groups_input',
                    placeholder="1-3\n5,6\n8-10",
                    height=100
                )
            
            st.divider()
            
            processing_disabled = (
                st.session_state.processing_status == 'in_progress' or
                not st.session_state.api_key
            )
            
            button_text = (
                "🔄 Przetwarzanie..."
                if st.session_state.processing_status == 'in_progress'
                else "🚀 Rozpocznij Przetwarzanie"
            )
            
            if st.button(
                button_text,
                use_container_width=True,
                type="primary",
                disabled=processing_disabled
            ):
                start_ai_processing()
                st.rerun()
            
            st.divider()
            
            st.info(f"**Projekt:** {st.session_state.project_name}")
            st.metric("Liczba stron", st.session_state.total_pages)
            st.caption(f"**Format:** {st.session_state.file_type.upper()}")

def render_processing_status():
    """Renderuje status przetwarzania"""
    if st.session_state.processing_status == 'idle' or not st.session_state.document:
        return
    
    processed_count = sum(1 for p in st.session_state.extracted_pages if p is not None)
    
    if st.session_state.processing_mode == 'article':
        total_groups = len(st.session_state.article_groups)
        processed_groups = st.session_state.next_article_index
        progress = processed_groups / total_groups if total_groups > 0 else 0
        
        if st.session_state.processing_status == 'complete':
            st.success(f"✅ Przetwarzanie zakończone! Przetworzono {total_groups} artykuł(ów).")
            
            # Przycisk szybkiej nawigacji do pierwszego artykułu
            if st.session_state.article_groups:
                if st.button("📖 Przejdź do pierwszego artykułu", type="secondary"):
                    st.session_state.current_page = st.session_state.article_groups[0][0] - 1
                    st.rerun()
        else:
            st.info(f"🔄 Przetwarzanie artykułów... ({processed_groups}/{total_groups})")
            st.progress(progress)
    else:
        progress = processed_count / st.session_state.total_pages if st.session_state.total_pages > 0 else 0
        
        if st.session_state.processing_status == 'complete':
            st.success("✅ Przetwarzanie zakończone!")
            
            # Przyciski szybkiej nawigacji do zakresu
            if st.session_state.processing_mode == 'range':
                nav_button_cols = st.columns(2)
                if nav_button_cols[0].button("📖 Przejdź do początku zakresu", type="secondary"):
                    st.session_state.current_page = st.session_state.start_page - 1
                    st.rerun()
                if nav_button_cols[1].button("📖 Przejdź do końca zakresu", type="secondary"):
                    st.session_state.current_page = st.session_state.end_page - 1
                    st.rerun()
        else:
            st.info(f"🔄 Przetwarzanie w toku... (Ukończono {processed_count}/{st.session_state.total_pages} stron)")
            st.progress(progress)
    
    c1, c2, _ = st.columns([1, 1, 3])
    
    if c1.button("💾 Zapisz postęp", use_container_width=True):
        save_project()
    
    articles = [
        p for p in st.session_state.extracted_pages
        if p and p.get('type') == 'artykuł' and p.get('is_group_lead', True)
    ]
    
    if articles:
        zip_data = [
            {
                'name': f"artykul_ze_str_{a['page_number']}.txt",
                'content': a['raw_markdown'].encode('utf-8')
            }
            for a in articles if 'raw_markdown' in a
        ]
        
        if zip_data:
            c2.download_button(
                "📥 Pobierz artykuły",
                create_zip_archive(zip_data),
                f"{st.session_state.project_name}_artykuly.zip",
                "application/zip",
                use_container_width=True
            )

def render_navigation():
    """Renderuje nawigację między stronami"""
    if st.session_state.total_pages <= 1:
        return
    
    st.subheader("📖 Nawigacja")
    
    # Informacja o zakresie przetwarzania
    if st.session_state.processing_mode == 'range':
        processing_range = f"{st.session_state.start_page}-{st.session_state.end_page}"
        st.info(f"🎯 Przetwarzany zakres: strony {processing_range}")
        
        # Przyciski szybkiej nawigacji
        nav_cols = st.columns(3)
        if nav_cols[0].button("⏮️ Początek zakresu", use_container_width=True):
            st.session_state.current_page = st.session_state.start_page - 1
            st.rerun()
        if nav_cols[1].button("⏭️ Koniec zakresu", use_container_width=True):
            st.session_state.current_page = st.session_state.end_page - 1
            st.rerun()
        if nav_cols[2].button("🏠 Początek dokumentu", use_container_width=True):
            st.session_state.current_page = 0
            st.rerun()
        
        st.divider()
    
    elif st.session_state.processing_mode == 'article' and st.session_state.article_groups:
        st.info(f"🎯 Liczba artykułów: {len(st.session_state.article_groups)}")
        
        # Przyciski nawigacji do artykułów
        article_nav_cols = st.columns(min(len(st.session_state.article_groups), 5))
        for idx, group in enumerate(st.session_state.article_groups[:5]):
            label = f"Art. {idx+1}"
            if len(group) > 1:
                label += f" ({group[0]}-{group[-1]})"
            else:
                label += f" (str. {group[0]})"
            
            if article_nav_cols[idx % 5].button(label, use_container_width=True, key=f"nav_art_{idx}"):
                st.session_state.current_page = group[0] - 1
                st.rerun()
        
        if len(st.session_state.article_groups) > 5:
            st.caption(f"... i jeszcze {len(st.session_state.article_groups) - 5} artykułów")
        
        st.divider()
    
    # Standardowa nawigacja
    c1, c2, c3 = st.columns([1, 2, 1])
    
    if c1.button(
        "⬅️ Poprzednia",
        use_container_width=True,
        disabled=(st.session_state.current_page == 0)
    ):
        st.session_state.current_page -= 1
        st.rerun()
    
    c2.metric("Strona", f"{st.session_state.current_page + 1} / {st.session_state.total_pages}")
    
    if c3.button(
        "Następna ➡️",
        use_container_width=True,
        disabled=(st.session_state.current_page >= st.session_state.total_pages - 1)
    ):
        st.session_state.current_page += 1
        st.rerun()
    
    # Slider nawigacji
    new_page = st.slider(
        "Przejdź do strony:",
        1,
        st.session_state.total_pages,
        st.session_state.current_page + 1
    ) - 1
    
    if new_page != st.session_state.current_page:
        st.session_state.current_page = new_page
        st.rerun()

def render_page_view():
    """Renderuje widok strony"""
    st.divider()
    
    page_index = st.session_state.current_page
    page_content = st.session_state.document.get_page_content(page_index)
    
    pdf_col, text_col = st.columns(2, gap="large")
    
    with pdf_col:
        st.subheader(f"📄 Oryginał (Strona {page_index + 1})")
        
        if st.session_state.file_type == 'pdf':
            image_data = st.session_state.document.render_page_as_image(page_index)
            if image_data:
                st.image(image_data, use_container_width=True)
            else:
                st.error("Nie można wyświetlić podglądu strony.")
        else:
            st.info(f"Podgląd nie jest dostępny dla plików {st.session_state.file_type.upper()}.")
        
        if page_content.images:
            with st.expander(f"🖼️ Pokaż/ukryj {len(page_content.images)} obraz(y)"):
                for img in page_content.images:
                    st.image(
                        img['image'],
                        caption=f"Obraz {img['index'] + 1}",
                        use_container_width=True
                    )
            
            img_zip = create_zip_archive([
                {
                    'name': f"str_{page_index+1}_img_{i['index']}.{i['ext']}",
                    'content': i['image']
                }
                for i in page_content.images
            ])
            
            st.download_button(
                "Pobierz obrazy",
                img_zip,
                f"obrazy_strona_{page_index+1}.zip",
                "application/zip",
                use_container_width=True
            )
    
    with text_col:
        st.subheader("🤖 Tekst przetworzony przez AI")
        
        with st.expander("👁️ Pokaż surowy tekst wejściowy"):
            st.text_area(
                "Surowy tekst",
                page_content.text,
                height=200,
                disabled=True,
                key=f"raw_text_{page_index}"
            )
        
        page_result = st.session_state.extracted_pages[page_index]
        
        if page_result:
            page_type = page_result.get('type', 'nieznany')
            color_map = {
                "artykuł": "green",
                "reklama": "orange",
                "pominięta": "grey",
                "błąd": "red"
            }
            color = color_map.get(page_type, "red")
            
            st.markdown(
                f"**Status:** <span style='color:{color}; text-transform:uppercase;'>"
                f"**{page_type}**</span>",
                unsafe_allow_html=True
            )
            
            group_pages = page_result.get('group_pages', [])
            if group_pages and len(group_pages) > 1:
                st.info(f"Ten artykuł obejmuje strony: {', '.join(str(p) for p in group_pages)}.")
            
            st.markdown(
                f"<div class='page-text-wrapper'>{page_result.get('formatted_content', '')}</div>",
                unsafe_allow_html=True
            )
            
            # --- PRZYCISKI AKCJI ---
            st.write("---")
            st.markdown("**Akcje Redakcyjne:**")
            
            allow_actions = (
                page_type == 'artykuł' and
                'raw_markdown' in page_result and
                page_result.get('is_group_lead', True)
            )

            action_cols = st.columns(4)
            
            if action_cols[0].button(
                "🔄 Przetwórz ponownie",
                key=f"reroll_{page_index}",
                use_container_width=True
            ):
                handle_page_reroll(page_index)
            
            if action_cols[1].button(
                "✨ Generuj Meta",
                key=f"meta_{page_index}",
                use_container_width=True,
                disabled=not allow_actions
            ):
                handle_meta_tag_generation(page_index, page_result['raw_markdown'])
            
            if action_cols[2].button(
                "🚀 Optymalizuj dla SEO",
                key=f"seo_{page_index}",
                use_container_width=True,
                disabled=not allow_actions,
                help="Przepisz artykuł zgodnie z zasadami SEO"
            ):
                handle_seo_generation(page_index, page_result['raw_markdown'])

            # Checkbox do pokazywania HTML
            show_html = action_cols[3].checkbox(
                "📄 Pokaż HTML",
                key=f"show_html_checkbox_{page_index}",
                disabled=not allow_actions,
                help="Pokaż i pobierz czysty HTML artykułu"
            )
            
            # WYŚWIETLENIE HTML JEŚLI CHECKBOX ZAZNACZONY
            if show_html and allow_actions:
                html_data = get_article_html_from_page(page_index)
                
                if html_data:
                    st.divider()
                    
                    with st.expander("📄 Czysty HTML artykułu", expanded=True):
                        st.caption(f"**{html_data['title']}**")
                        
                        # Taby dla różnych widoków
                        tab1, tab2 = st.tabs(["💻 Kod HTML (zawartość)", "📰 Pełny dokument HTML"])
                        
                        with tab1:
                            st.code(html_data['html_content'], language='html', line_numbers=True)
                            
                            st.download_button(
                                label="📥 Pobierz zawartość HTML",
                                data=html_data['html_content'],
                                file_name=f"{sanitize_filename(html_data['title'])}_content.html",
                                mime="text/html",
                                use_container_width=True,
                                key=f"download_content_{page_index}"
                            )
                        
                        with tab2:
                            st.code(html_data['html_document'], language='html', line_numbers=True)
                            
                            st.download_button(
                                label="📥 Pobierz pełny dokument HTML",
                                data=html_data['html_document'],
                                file_name=f"{sanitize_filename(html_data['title'])}.html",
                                mime="text/html",
                                use_container_width=True,
                                key=f"download_full_{page_index}"
                            )
                        
                        if html_data['meta_title'] or html_data['meta_description']:
                            st.info("ℹ️ Ten HTML zawiera wygenerowane meta tagi SEO")
            
            # --- WYNIKI AKCJI ---

            # META TAGI
            if page_index in st.session_state.meta_tags:
                tags = st.session_state.meta_tags[page_index]
                
                if "error" in tags:
                    st.error(f"Błąd generowania meta tagów: {tags['error']}")
                else:
                    with st.expander("Wygenerowane Meta Tagi ✨", expanded=False):
                        st.text_input(
                            "Meta Title",
                            value=tags.get("meta_title", ""),
                            key=f"mt_{page_index}"
                        )
                        st.text_area(
                            "Meta Description",
                            value=tags.get("meta_description", ""),
                            key=f"md_{page_index}"
                        )
            
            if page_index in st.session_state.seo_articles:
                seo_result = st.session_state.seo_articles[page_index]
                with st.expander("🤖 Zoptymalizowany Artykuł SEO", expanded=True):
                    if "error" in seo_result:
                        st.error(f"Błąd podczas optymalizacji SEO: {seo_result['error']}")
                        st.json(seo_result)
                    else:
                        seo_title = seo_result.get("seo_title", "Brak tytułu")
                        seo_markdown = seo_result.get("seo_article_markdown", "Brak treści.")
                        
                        st.markdown(f"### {seo_title}")
                        st.markdown(seo_markdown, unsafe_allow_html=True)
                        st.download_button(
                            label="📥 Pobierz wersję SEO (.txt)",
                            data=f"# {seo_title}\n\n{seo_markdown}",
                            file_name=f"{sanitize_filename(seo_title)}.txt",
                            mime="text/plain",
                            use_container_width=True,
                            key=f"download_seo_{page_index}"
                        )

        else:
            if st.session_state.processing_status == 'in_progress':
                st.info("⏳ Strona oczekuje na przetworzenie...")
            else:
                st.info("Uruchom przetwarzanie w panelu bocznym.")

def handle_page_reroll(page_index: int):
    """Przetwarza stronę ponownie z kontekstem"""
    with st.spinner("Przetwarzanie strony z kontekstem..."):
        prev_text = ""
        if page_index > 0:
            prev_content = st.session_state.document.get_page_content(page_index - 1)
            prev_text = prev_content.text
        
        curr_content = st.session_state.document.get_page_content(page_index)
        curr_text = curr_content.text
        
        next_text = ""
        if page_index < st.session_state.total_pages - 1:
            next_content = st.session_state.document.get_page_content(page_index + 1)
            next_text = next_content.text
        
        context_text = (
            f"KONTEKST (POPRZEDNIA STRONA):\n{prev_text}\n\n"
            f"--- STRONA DOCELOWA ---\n{curr_text}\n\n"
            f"KONTEKST (NASTĘPNA STRONA):\n{next_text}"
        )
        
        ai_processor = AIProcessor(st.session_state.api_key, st.session_state.model)
        page_content = PageContent(page_index + 1, context_text)
        new_result = asyncio.run(ai_processor.process_page(page_content))
        
        st.session_state.extracted_pages[page_index] = new_result
    
    st.rerun()

def handle_meta_tag_generation(page_index: int, raw_markdown: str):
    """Generuje meta tagi dla artykułu"""
    with st.spinner("Generowanie meta tagów..."):
        ai_processor = AIProcessor(st.session_state.api_key, st.session_state.model)
        tags = asyncio.run(ai_processor.generate_meta_tags(raw_markdown))
        st.session_state.meta_tags[page_index] = tags
    
    st.rerun()

def handle_seo_generation(page_index: int, raw_markdown: str):
    """Generuje zoptymalizowaną wersję artykułu"""
    with st.spinner("🚀 Optymalizowanie artykułu dla SEO... To może chwilę potrwać."):
        ai_processor = AIProcessor(st.session_state.api_key, st.session_state.model)
        result = asyncio.run(ai_processor.generate_seo_article(raw_markdown))
        st.session_state.seo_articles[page_index] = result
    
    st.rerun()

# ===== GŁÓWNA APLIKACJA =====

def main():
    st.set_page_config(
        layout="wide",
        page_title="Redaktor AI - Procesor Dokumentów",
        page_icon="🚀"
    )
    
    st.markdown("""
    <style>
    .page-text-wrapper {
        border: 1px solid #e0e0e0;
        border-radius: 8px;
        padding: 20px;
        background-color: #f9f9f9;
        max-height: 600px;
        overflow-y: auto;
    }
    .error-box {
        background-color: #ffebee;
        border-left: 4px solid #f44336;
        padding: 12px;
        border-radius: 4px;
        margin: 10px 0;
    }
    .stButton button {
        border-radius: 8px;
    }
    h2, h3, h4 {
        margin-top: 1em;
        margin-bottom: 0.5em;
    }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("🚀 Redaktor AI - Interaktywny Procesor Dokumentów")
    
    init_session_state()
    
    if not DOCX_AVAILABLE or not MAMMOTH_AVAILABLE:
        missing = []
        if not DOCX_AVAILABLE:
            missing.append("DOCX (zainstaluj: pip install python-docx)")
        if not MAMMOTH_AVAILABLE:
            missing.append("DOC (zainstaluj: pip install mammoth)")
        
        with st.sidebar:
            with st.expander("⚠️ Ograniczona funkcjonalność", expanded=False):
                st.warning("Niektóre formaty plików nie są dostępne:")
                for fmt in missing:
                    st.write(f"- {fmt}")
    
    if not st.session_state.api_key:
        st.error("❌ Brak klucza API OpenAI!")
        st.info("Proszę skonfiguruj swój klucz API w Streamlit Secrets.")
        st.stop()
    
    render_sidebar()
    
    if not st.session_state.document:
        if not st.session_state.project_loaded_and_waiting_for_file:
            st.info("👋 Witaj! Aby rozpocząć, wgraj plik (PDF/DOCX/DOC) lub załaduj istniejący projekt z panelu bocznego.")
            
            with st.expander("📖 Jak korzystać z aplikacji?"):
                st.markdown("""
                ### Tryby przetwarzania:
                
                1. **Cały dokument** - Przetwarza każdą stronę osobno, jedno zapytanie na stronę
                2. **Zakres stron** - Przetwarza wybrany zakres stron osobno
                3. **Artykuł wielostronicowy** - Łączy wybrane strony i przetwarza jako jeden artykuł w jednym zapytaniu
                
                ### Obsługiwane formaty:
                - PDF (z podglądem i wyciąganiem grafik)
                - DOCX (Microsoft Word)
                - DOC (starsze pliki Word)
                
                ### Funkcje:
                - Zapisywanie i ładowanie projektów
                - Wyciąganie grafik ze stron
                - Generowanie meta tagów SEO
                - **Nowość! Optymalizacja artykułów pod kątem SEO**
                - **Eksport do HTML** - czysty HTML z obsługą list i przypisów
                - Ponowne przetwarzanie stron z kontekstem
                """)
        return
    
    render_processing_status()
    
    if st.session_state.processing_status == 'in_progress':
        run_ai_processing_loop()
    else:
        render_navigation()
        render_page_view()

if __name__ == "__main__":
    main()
